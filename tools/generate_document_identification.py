from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = "document_identification.docx"
SOFTWARE = "可穿戴双向手语-语音交互固件系统"
VERSION = "V1.0"
HEADER = f"{SOFTWARE} {VERSION}"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    set_run_font(r, "宋体", 10.5)


def set_run_font(run, font="宋体", size=10.5):
    run.font.name = font
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def add_field(paragraph, instr):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(text)
    run._r.append(fld_char3)
    set_run_font(run)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    p.add_run("第 ")
    add_field(p, "PAGE")
    p.add_run(" 页 / 共 ")
    add_field(p, "SECTIONPAGES")
    p.add_run(" 页")
    for r in p.runs:
        set_run_font(r, "宋体", 9)


def configure_section(section, restart=False, with_footer=True):
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.text = HEADER
    for r in hp.runs:
        set_run_font(r, "宋体", 9)
    if restart:
        sect_pr = section._sectPr
        pg_num = sect_pr.find(qn("w:pgNumType"))
        if pg_num is None:
            pg_num = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num)
        pg_num.set(qn("w:start"), "1")
    if with_footer:
        add_page_number_footer(section)
    else:
        section.footer.paragraphs[0].text = ""


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(0)
    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        st = doc.styles[name]
        st.font.name = "宋体"
        st.font.size = Pt(size)
        st.font.bold = True
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        st.paragraph_format.space_before = Pt(3)
        st.paragraph_format.space_after = Pt(2)
        st.paragraph_format.line_spacing = 1.2
    if "CodeBlock" not in [s.name for s in doc.styles]:
        code = doc.styles.add_style("CodeBlock", 1)
    else:
        code = doc.styles["CodeBlock"]
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.space_after = Pt(0)


def add_lines(doc, lines, code=False):
    style = "CodeBlock" if code else "Normal"
    for line in lines:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.line_spacing = 1.0 if code else 1.2
        r = p.add_run(line)
        set_run_font(r, "Consolas" if code else "宋体", 9 if code else 10.5)


def add_page(doc, title, lines, heading=1):
    if title:
        doc.add_heading(title, level=heading)
    add_lines(doc, lines)
    doc.add_page_break()


def pad(lines, target=50, prefix="说明"):
    return list(lines)


def bullet(title, items):
    return [title] + [f"{i+1}. {v}" for i, v in enumerate(items)]


def code_page(doc, title, intro, code_lines):
    doc.add_heading(title, level=2)
    add_lines(doc, pad(intro, 16, "代码说明"))
    add_lines(doc, ["```c"] + code_lines + ["```"], code=True)
    add_lines(doc, pad(["以上片段为源码关键路径摘录式说明，不替代完整源程序。"], 18, "实现说明"))
    doc.add_page_break()


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SOFTWARE)
    r.bold = True
    set_run_font(r, "宋体", 22)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(VERSION)
    r.bold = True
    set_run_font(r, "宋体", 18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("软件说明书")
    r.bold = True
    set_run_font(r, "宋体", 18)
    for _ in range(8):
        doc.add_paragraph("")
    for line in ["软件名称：可穿戴双向手语-语音交互固件系统", "版本号：V1.0", "文档类型：软件说明书", "编写日期：【待填写】", "著作权人：【待填写】"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, "宋体", 12)
    doc.add_page_break()


def add_revision(doc):
    doc.add_heading("修订历史", level=1)
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    headers = ["版本", "日期", "修订内容", "编写", "审核"]
    vals = ["V1.0", "【待填写】", "形成软件著作权申报用软件说明书初稿。", "【待填写】", "【待填写】"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_text(table.rows[1].cells[i], vals[i])
    add_lines(doc, pad([
        "文档用途：作为中国版权保护中心软件著作权登记的文档鉴别材料。",
        "适用软件：可穿戴双向手语-语音交互固件系统 V1.0。",
        "说明范围：嵌入式固件工程，包含左手端、右手端及其 BLE 通信协议；上位机识别程序不在本次申报范围内。",
        "保密提示：涉及 API Key、APISecret、无线网络口令等敏感配置时，应在提交版本中脱敏。"
    ], 44, "修订说明"))
    doc.add_page_break()


def add_toc(doc):
    lines = [
        "目录",
        "第 1 章 概述 ........................................................................ 正文 P1",
        "第 2 章 系统架构 .................................................................... 正文 P5",
        "第 3 章 功能模块详述 ................................................................ 正文 P9",
        "模块 1：系统启动与初始化模块 .................................................. 正文 P9",
        "模块 2：IMU 数据采集模块 ....................................................... 正文 P14",
        "模块 3：BLE 无线通信模块 ....................................................... 正文 P18",
        "模块 4：BLE CSV 分片与帧序列模块 ........................................... 正文 P24",
        "模块 5：运行模式管理模块 ....................................................... 正文 P30",
        "模块 8：OLED 状态显示模块 ..................................................... 正文 P34",
        "模块 9：电池与电源监测模块 ..................................................... 正文 P37",
        "模块 10：辅助通信与调试模块 .................................................... 正文 P39",
        "模块 11：硬件接口与引脚分配表 .................................................. 正文 P42",
        "模块 12：术语与参考文档 ......................................................... 正文 P44",
        "附录 A：测试与验证 ............................................................... 正文 P46",
        "附录 B：配置与部署 ............................................................... 正文 P49",
        "模块 6：左手语音播报模块（TTS） ............................................. 正文 P52",
        "模块 7：右手音频采集与 AI 云识别模块（STT） ............................... 正文 P55",
        "待确认事项 .......................................................................... 正文 P60",
    ]
    add_lines(doc, pad(lines, 50, "目录说明"))
    doc.add_page_break()
    add_lines(doc, pad([
        "目录说明：本说明书采用显式分页编排，便于打印前 30 页和后 30 页时展示核心模块。",
        "前 30 页覆盖模块 1 至模块 4，重点展示系统初始化、IMU 采集、BLE 服务和自定义分片协议。",
        "后 30 页覆盖模块 6 和模块 7，重点展示左手 TTS 与右手 STT 的双向交互特性。",
    ], 50, "目录补充"))


def main():
    doc = Document()
    setup_styles(doc)
    configure_section(doc.sections[0], restart=False, with_footer=False)
    add_cover(doc)
    add_revision(doc)
    add_toc(doc)

    body = doc.add_section(WD_SECTION_START.NEW_PAGE)
    body.header.is_linked_to_previous = False
    body.footer.is_linked_to_previous = False
    configure_section(body, restart=True, with_footer=True)

    add_page(doc, "第 1 章 概述", pad([
        "1.1 软件名称与版本",
        f"软件名称：{SOFTWARE}。",
        "版本号：V1.0。",
        "软件类型：嵌入式固件系统。",
        "运行载体：ART-Pi2 开发板，主控为 STM32H7R7 系列 MCU，运行 RT-Thread RTOS。",
        "左手端设备名：ART-Pi2-IMU-L；右手端设备名：ART-Pi2-IMU-R。",
        "BLE 配置文档记录的左手 MAC 地址为 C0:4E:51:05:34:33，右手 MAC 地址为 C0:31:51:05:34:33；该地址由 CYW43438 蓝牙芯片固件提供，应用层不直接配置。",
        "RT-Thread 配置启用组件初始化、用户 main、finsh/MSH、串口、I2C、SDIO、WiFi、SAL、netdev、lwIP 2.1.2、FAL、DFS 和 POSIX socket。",
        "无线驱动侧启用 WHD_USING_CHIP_CYW43438 与 WHD_USING_WIFI5，WiFi 通过 SDIO/WHD 工作，蓝牙通过 UART7 HCI H4 通道接入 NimBLE Host。",
        "固件以 C 语言实现，采用 Keil MDK-ARM、SCons 与 RT-Thread ENV 进行构建。",
        "源程序统计：自写源代码约 18,702 行，去空行注释约 11,680 行，自写 .c/.h 文件约 85 个。",
        "左手工程约 35 个自写文件，右手工程约 50 个自写文件。",
    ], 50, "概述"))
    add_page(doc, "1.2 开发目的与背景", pad([
        "本软件面向听障人士与健听人群之间的双向交流辅助需求。",
        "系统通过双手多路 IMU 采集手势运动信息，并通过 BLE 将结构化数据发送给上位机识别程序。",
        "左手端具备 VTX316 语音合成功能，可把上位机或对端传入的文本朗读出来。",
        "右手端具备 INMP441 音频采集与云端语音识别链路，可将现场语音转写为文字并显示。",
        "上述两条链路共同构成手语到语音、语音到文字的双向辅助交互闭环。",
        "本次申报范围为运行在 ART-Pi2 上的嵌入式固件，不包含 ROCK 5B 上位机识别程序。",
        "工程中同时保留 WiFi TCP 辅助上传链路，默认服务器地址为 192.168.221.92:8266；该链路用于调试、数据显示和翻译文本转发，不替代 BLE 申报核心协议。",
        "右手 STT 的讯飞接入在固件侧通过 HTTP 代理 URL 调用，真实 WebSocket/HMAC/IAT 细节由 PC 代理脚本承接，提交材料中已将其边界列入待确认。",
    ], 50, "背景"))
    add_page(doc, "1.3 面向用户与应用场景", pad([
        "面向用户：听障人士、助残辅具开发者、康复训练场景使用者、手语交互研究人员。",
        "应用场景一：双手穿戴设备采集手势数据，供上位机进行手语识别。",
        "应用场景二：健听人群语音经右手端 STT 链路转写后，用 OLED 或后续链路显示给听障用户。",
        "应用场景三：外部识别结果通过 BLE Text 特征写入左手端，由 VTX316 模块播报。",
        "应用场景四：多模态感知与人机交互实验平台，验证 IMU、BLE、音频和显示的协同工作。",
    ], 50, "场景"))
    add_page(doc, "1.4 软件特点摘要", pad([
        "特点一：双手 11 路 IMU 同步采集，包含 10 路 MPU6050 与 1 路 ICM-20948。",
        "特点二：基于 NimBLE 的 BLE GATT 自定义服务，服务 UUID 为 A74D0001-B4E7-4C5F-9D2A-F163E80ACB00。",
        "特点三：自定义 [FRAG] 文本分片协议，用于突破 BLE MTU 对长 CSV 帧的限制。",
        "特点四：AUTO/MANUAL 与 STANDBY/SLEEP/RUNNING 组合状态机，支持本地 PC6 按键与 BLE CMD 远程控制。",
        "特点五：左手端集成 VTX316 TTS，右手端集成 INMP441 采集与讯飞识别代理链路。",
        "特点六：OLED 状态显示、电池 ADC 监测、分级日志与 finsh shell 调试接口便于现场维护。",
        "特点七：主线程包含启动自检与错误提示路径，左手可通过 VTX316 播报传感器、WiFi、TCP、语音和 BLE 启动异常，右手可通过 OLED 显示 I2C、UART、IMU、WiFi、TCP 和 BLE 异常。",
        "特点八：I2C2 与 OLED/IMU 共享场景下提供 i2c2_mutex_take/release 互斥接口，降低 TCA9548A 通道切换与 OLED 刷新之间的总线冲突风险。",
    ], 50, "特点"))

    add_page(doc, "第 2 章 系统架构", pad([
        "2.1 整体硬件架构",
        "系统由左手穿戴端、右手穿戴端和 ROCK 5B 上位机三部分组成。",
        "左手端与右手端均以 ART-Pi2 为主控平台，包含 STM32H7R7 MCU 与 CYW43438 WiFi/BT Combo 模组。",
        "主控启动时通过 INIT_BOARD_EXPORT 执行 vtor_config，将中断向量表基址重定位到 XSPI2_BASE，以适配 ART-Pi2 外部 XSPI2 Flash 启动方式。",
        "左手 main.c 负责 IIC、VTX316、MPU、battery 线程创建，并在 WiFi 连接成功后自动启动 TCP；右手 main.c 额外创建 uart_oled 与 autostart 线程。",
        "右手 autostart 线程按 WiFi 连接、网络就绪、传感器校准、TCP 客户端、voice_assistant 的顺序执行一次性启动编排。",
        "左右手共有外设包括 10 路 MPU6050、1 路 ICM-20948、SSD1306 OLED、PC6 按键和锂电池 ADC 采集。",
        "左手独有外设为 VTX316 语音合成模块；右手独有外设为 INMP441 数字麦克风。",
        "[此处插入系统架构图]",
    ], 50, "架构"))
    add_page(doc, "2.2 软件分层架构", pad([
        "应用层：主线程、BLE 应用、IMU Notify 线程、运行模式状态机、OLED 显示、TTS/STT 业务任务。",
        "协议层：NimBLE Host、GATT 服务、[DATA]/[FRAG] 文本帧、CMD 命令解析、HTTP 代理访问。",
        "网络层：RT-Thread WLAN 管理、SAL/POSIX socket、lwIP 2.1.2、TCP 客户端、HTTP web_client。",
        "资源层：WHD 固件、CLM、NVRAM 资源通过 FAL 外部存储分区加载，rtconfig.h 中配置 WHD_RESOURCES_IN_EXTERNAL_STORAGE_FAL。",
        "驱动层：CYW43438 HCI 适配器、I2C/TCA9548A、MPU6050、ICM-20948、SAI/INMP441、UART、ADC。",
        "操作系统层：RT-Thread 线程调度、信号量、互斥量、初始化导出机制、finsh shell 与日志组件。",
        "构建层：Keil MDK-ARM μVision5、ARM Compiler V6.22、SCons 与 RT-Thread ENV。",
    ], 50, "分层"))
    add_page(doc, "2.3 左右手差异化外设方案对比表", pad([
        "左手端：设备名 ART-Pi2-IMU-L，主要差异外设为 VTX316 TTS 模块。",
        "左手端链路：BLE Text 写入后，经 _ble_text_recv_handler 规范化文本，再调用 vtx316_speak_wait 同步播报。",
        "右手端：设备名 ART-Pi2-IMU-R，主要差异外设为 INMP441 数字麦克风。",
        "右手端链路：INMP441 经 SAI/DMA 采集音频，audio_process 执行降噪、能量计算和 VAD，voice_assistant 调用 ai_cloud_service_speech_to_text。",
        "共有链路：左右手均注册相同 IMU GATT 服务，均通过 imu_notify_thread 周期推送 11 路 IMU CSV 分片。",
        "左手 SConscript 为 applications 额外加入 NimBLE v1.0.0 的 host、services、store、porting 和 npl/rtthread 头文件路径；右手应用目录则以本地应用源文件为主。",
        "左手 vtx316 组件依赖 BSP_USING_UART1，保证 UART1 驱动开启后才编译语音合成模块。",
    ], 50, "对比"))
    add_page(doc, "2.4 数据流与控制流示意", pad([
        "IMU 数据流：传感器阵列 -> I2C 多路复用 -> mpu_get_channel_raw_data -> CSV 组帧 -> [FRAG] 分片 -> BLE Notify -> 上位机。",
        "TTS 数据流：BLE Text 写入 -> CMD 前缀拦截 -> 文本规范化 -> VTX316 串口协议 -> 语音播报。",
        "STT 数据流：INMP441 -> SAI DMA -> audio_capture/audio_process -> VAD 端点检测 -> ai_cloud_service -> OLED 显示。",
        "控制流：PC6 按键或 BLE CMD 触发运行模式切换，operation_mode_ble_notify_enabled 控制 IMU Notify 是否发送。",
        "BLE 与 TCP 的控制边界不同：operation_mode 状态机只控制 BLE IMU Notify 与 frame_seq 复位，不改变 TCP 上传行为。",
        "文本下行边界不同：BLE Text 特征用于 CMD 或左手 TTS；TCP 接收回调也支持 SAY: 文本，但属于辅助链路。",
        "[此处插入数据流与控制流示意图]",
    ], 50, "数据流"))

    # Module 1: 5 pages
    add_page(doc, "第 3 章 功能模块详述", pad([
        "模块 1：系统启动与初始化模块",
        "对应源文件：main.c、ble_app.c 初始化段。",
        "RT-Thread 初始化流程按 BOARD、DEVICE、COMPONENT、ENV、APP 等阶段执行。",
        "CYW43438 蓝牙控制器初始化位于 DEVICE 阶段，NimBLE port 初始化位于 COMPONENT 阶段，_ble_app_init 位于 ENV 阶段。",
        "CYW43438 BT 传输层使用 uart7，默认 HCI 波特率 115200，初始化时打开 UART7、发送 HCI Reset，并在 Reset Complete 到达后设置 ready 标志。",
        "NimBLE Host 线程配置来源于 syscfg.h，主机线程栈 4096 字节，优先级 17，最大连接数为 1，ATT Preferred MTU 为 247。",
        "主业务线程在 main.c 中创建，包括 IIC/OLED、UART/OLED、MPU6050/ICM-20948、battery 等线程。",
        "[此处插入启动时序图]",
    ], 50, "启动"))
    add_page(doc, "模块 1：_ble_app_init 初始化流程", pad([
        "第一步：operation_mode_init 初始化运行模式状态机，启动默认状态 AUTO / AUTO_STANDBY。",
        "第二步：等待 cyw43438_bt_is_ready 返回真，最长等待 100 次，每次 50 ms。",
        "第三步：调用 nimble_hci_adapter_init 安装 HCI transport shim。",
        "第四步：配置 ble_hs_cfg.sync_cb 与 ble_hs_cfg.reset_cb。",
        "第五步：初始化标准 GAP/GATT 服务，并设置设备名 ART-Pi2-IMU-L 或 ART-Pi2-IMU-R。",
        "第六步：调用 imu_ble_service_init 注册自定义 IMU GATT 服务。",
        "第七步：启动 NimBLE Host 线程和 IMU Notify 线程。",
        "左手端另注册 imu_ble_text_set_callback(_ble_text_recv_handler)，用于 VTX316 语音播报。",
        "左手初始化流程中第 6 步注册 VTX316 文本回调，右手没有该回调；两端均在第 7/8 步启动 Host 与 Notify 线程。",
    ], 50, "初始化"))
    code_page(doc, "模块 1：初始化关键代码片段", [
        "以下代码片段概括 _ble_app_init 的主要步骤，左手端和右手端流程基本一致。",
    ], [
        "static int _ble_app_init(void)",
        "{",
        "    operation_mode_init();",
        "    while (!cyw43438_bt_is_ready() && waited < 100) {",
        "        rt_thread_mdelay(50);",
        "        waited++;",
        "    }",
        "    if (!cyw43438_bt_is_ready()) return -RT_ETIMEOUT;",
        "    ret = nimble_hci_adapter_init();",
        "    ble_hs_cfg.sync_cb  = _on_sync;",
        "    ble_hs_cfg.reset_cb = _on_reset;",
        "    ble_svc_gap_init();",
        "    ble_svc_gatt_init();",
        "    ble_svc_gap_device_name_set(DEVICE_NAME);",
        "    ret = imu_ble_service_init();",
        "    ble_hs_thread_startup();",
        "    return imu_notify_thread_start();",
        "}",
    ])
    add_page(doc, "模块 1：失败路径与初始化标志", pad([
        "左手端使用 g_left_ble_init_ok 记录 BLE 应用初始化状态，右手端使用 g_right_ble_init_ok。",
        "若 CYW43438 未就绪、HCI 适配器初始化失败、GATT 服务注册失败或 Notify 线程启动失败，对应标志置为 RT_FALSE。",
        "当 NimBLE Host 完成同步后，_on_sync 推断地址类型、打印 BLE 地址并启动广播。",
        "连接失败或断开连接时，GAP 事件处理函数会重新调用 _start_advertising。",
        "初始化标志通过 left_ble_init_ok 或 right_ble_init_ok 向其他模块提供只读状态。",
        "左手主线程在初始化末尾检查 left_ble_init_ok，异常时调用 vtx316_report_boot_error(VTX316_BOOT_ERROR_BLE)。",
        "右手主线程检查 right_ble_init_ok，异常时调用 right_show_failure(\"BLE FAIL\") 显示 OLED 错误信息。",
    ], 50, "失败路径"))
    add_page(doc, "模块 1：主线程与并发关系", pad([
        "main.c 负责创建常驻业务线程和一次性初始化线程。",
        "IIC/OLED 线程负责 I2C 总线、TCA9548A 与 OLED 初始化。",
        "左手 iic_drv、vtx316 为一次性初始化线程，mpu6050、bat_adc 为常驻线程；右手 iic_drv 为一次性初始化线程，uart_oled、mpu6050、bat_adc 为常驻线程。",
        "左手在等待 WLAN 设备最多 10 秒、WiFi 连接最多 15 秒后，延迟 2 秒启动 TCP 客户端；右手 autostart 线程最多重试 WiFi 3 次。",
        "两端主循环均以 PO5 LED 进行 500 ms 亮、500 ms 灭的心跳指示。",
        "MPU 线程负责 MPU6050/ICM-20948 数据采样与缓存。",
        "battery 线程负责 ADC 电池采样与百分比映射。",
        "BLE Host 线程与 imu_ntfy 线程并行运行，前者处理协议栈事件，后者按 90 ms 周期尝试发送 IMU 分片。",
        "线程间共享状态采用 RT-Thread 原语、volatile 变量和中断临界区保护。",
    ], 50, "并发"))

    # Module 2: 4 pages
    add_page(doc, "模块 2：IMU 数据采集模块", pad([
        "对应源文件：mpu6050/*.c/.h、IIC/*.c/.h。",
        "硬件布局：10 路 MPU6050 采集手指关节六轴数据，1 路 ICM-20948 位于手背采集九轴数据。",
        "mpu_channel_data_t 包含 ax、ay、az、gx、gy、gz、mx、my、mz、valid 等字段。",
        "mpu_channel_data_t 还包含 has_mag，用于标记通道是否具备磁力计；仅 ch10 ICM-20948 正常情况下具有磁力计数据。",
        "MPU_TOTAL_CHANNELS = MPU6050_I2C1_COUNT + MPU6050_I2C2_COUNT，即 8 + 3 = 11。",
        "MPU6050 采集线程优先级为 16，栈大小为 4096 字节，时间片为 10。",
        "mpu_get_channel_raw_data(ch, out) 用于读取指定通道最新原始数据。",
        "ch0 至 ch9 输出 6 个整数，ch10 输出 9 个整数，CSV 总计 69 个 IMU 整数字段。",
        "[此处插入 IMU 阵列布局图]",
    ], 50, "IMU"))
    add_page(doc, "模块 2：I2C 多路复用与通道切换", pad([
        "系统通过 TCA9548A 等 I2C 多路复用器扩展多个 IMU 通道。",
        "I2C1 使用 PB8/PB9，连接 TCA9548A #1，ch0 至 ch7 对应 8 路 MPU6050。",
        "I2C2 使用 PE1/PE2，连接 TCA9548A #2，ch0、ch1 对应 2 路 MPU6050，ch2 对应 ICM-20948，ch3 用于 OLED。",
        "TCA9548A 地址为 0x70，通道数为 8，驱动提供 tca9548a_init、select_channel、disable_all_channels、get_current_channel 和 is_present。",
        "源码注释显示 I2C2 使用 PE1/PE2，并与 OLED 共享总线资源。",
        "通道切换后访问对应传感器地址，完成 WHO_AM_I 检测、寄存器配置与数据读取。",
        "I2C2 的通道 2 用于 ICM-20948，并启用 bypass 模式访问其内部 AK09916 磁力计。",
        "总线共享场景下通过互斥或串行化访问避免 OLED 与传感器读写冲突，具体实现详见源代码。",
        "i2c2_mutex_init 在右手 main.c 启动早期调用，所有 I2C2 共享访问可通过 i2c2_mutex_take(timeout_ms) 与 i2c2_mutex_release 包裹。",
    ], 50, "I2C"))
    add_page(doc, "模块 2：MPU6050 与 ICM-20948 初始化", pad([
        "MPU6050 初始化内容包括时钟源选择、量程配置、DLPF 配置和零偏校准。",
        "采集线程包含零偏校准样本数差异：ICM-20948 使用更多样本，并使用更强低通滤波系数，以降低手背九轴数据抖动。",
        "BLE Notify 路径明确读取 mpu_get_channel_raw_data，保持原始整数语义；TCP 辅助路径可使用经滤波/零偏处理后的数据，二者语义不同。",
        "ICM-20948 初始化内容包括 Bank 选择、WHO_AM_I 校验、PWR_MGMT_1 复位与时钟配置。",
        "ICM-20948 通过 INT_PIN_CFG 启用 I2C bypass 后访问 AK09916。",
        "AK09916 初始化包含 WIA2 身份校验、软复位和连续测量模式配置。",
        "ICM-20948 的陀螺仪与加速度计寄存器地址不同于 MPU6050，因此源码中实现了专用读取路径。",
    ], 50, "初始化"))
    code_page(doc, "模块 2：异常通道零值替代策略", [
        "BLE 组帧时若通道读取失败或 valid 为假，输出固定数量的 0，保证上位机解析字段数稳定。",
    ], [
        "mpu_channel_data_t d;",
        "if (mpu_get_channel_raw_data(i, &d) == RT_EOK && d.valid) {",
        "    pos += rt_snprintf(csv_buf + pos, remain,",
        "                      \",%d,%d,%d,%d,%d,%d\",",
        "                      d.ax, d.ay, d.az, d.gx, d.gy, d.gz);",
        "} else {",
        "    pos += rt_snprintf(csv_buf + pos, remain,",
        "                      \",0,0,0,0,0,0\");",
        "}",
    ])

    # Module 3: 6 pages
    for title, lines in [
        ("模块 3：BLE 无线通信模块", [
            "对应源文件：imu_ble_service.c/.h、ble_app.c GAP 事件。",
            "本模块基于 NimBLE V1.0.0，在 RT-Thread 上以 Host 线程方式运行。",
            "CYW43438 控制器通过 HCI 适配器与 NimBLE Host 连接。",
            "syscfg.h 中 NIMBLE_CFG_CONTROLLER = 0，表示本工程只运行 BLE Host，控制器位于 CYW43438 芯片内部。",
            "BLE ACL buffer 数量为 4，ACL buffer size 为 255，msys block size 为 292，block count 为 12。",
            "自定义 GATT 服务用于暴露 IMU Notify、通道选择和文本写入特征。",
            "[此处插入 GATT 服务表]",
        ]),
        ("模块 3：自定义 GATT 服务与特征", [
            "服务 UUID：A74D0001-B4E7-4C5F-9D2A-F163E80ACB00。",
            "Notify 特征 UUID：A74D0002-B4E7-4C5F-9D2A-F163E80ACB00，属性为 Read + Notify。",
            "Channel 特征 UUID：A74D0003-B4E7-4C5F-9D2A-F163E80ACB00，属性为 Read + Write，当前为通道选择保留。",
            "Text 特征 UUID：A74D0004-B4E7-4C5F-9D2A-F163E80ACB00，属性为 Write，用于文本或 CMD 命令。",
            "GATT 表由 ble_gatts_count_cfg 与 ble_gatts_add_svcs 注册。",
        ]),
        ("模块 3：广播与连接参数", [
            "广播模式为 BLE_GAP_CONN_MODE_UND，可连接非定向广播。",
            "发现模式为 BLE_GAP_DISC_MODE_GEN。",
            "广播间隔下限为 200 ms，上限为 500 ms，通过 BLE_GAP_ADV_ITVL_MS 转换为 0.625 ms 单位。",
            "左手广播名 ART-Pi2-IMU-L，右手广播名 ART-Pi2-IMU-R。",
            "代码中的 128-bit UUID 数组采用 LSB-first 字节序书写，文档中统一采用标准显示形式 A74D000x-B4E7-4C5F-9D2A-F163E80ACB00。",
            "BLE 配置文档给出左手 MAC C0:4E:51:05:34:33、右手 MAC C0:31:51:05:34:33，用于 ROCK 侧区分连接目标。",
            "[此处插入 BLE 连接流程图]",
        ]),
        ("模块 3：GAP 事件表", [
            "CONNECT：连接成功时保存 conn_handle，连接失败时清空句柄并重启广播。",
            "DISCONNECT：清除订阅状态和连接句柄，随后重启广播。",
            "SUBSCRIBE：根据 cur_notify 更新 imu_notify_set_subscribed。",
            "MTU：记录 conn、cid、mtu 值，用于调试 MTU 协商。",
            "CONN_UPDATE：连接参数更新事件当前返回 0，具体处理详见源代码。",
        ]),
        ("模块 3：CMD 前缀拦截机制", [
            "Text 特征写入时先复制 OS mbuf 数据到本地缓冲区。",
            "若文本以 CMD: 开头，则调用 operation_mode_handle_cmd 处理，不再转发到语音或翻译文本路径。",
            "当前支持 CMD:START、CMD:STOP、CMD:RESET_SEQ。",
            "未知 CMD 命令仅记录警告并忽略。",
            "非 CMD 文本由 imu_ble_text_set_callback 注册的回调继续处理。",
        ]),
    ]:
        add_page(doc, title, pad(lines, 50, "BLE"))
    code_page(doc, "模块 3：Text 特征访问路径片段", [
        "以下片段体现 CMD 前缀截获和普通文本回调分流。",
    ], [
        "if (strncmp(buf, \"CMD:\", 4) == 0) {",
        "    if (operation_mode_handle_cmd(buf) != RT_TRUE) {",
        "        LOG_W(\"unknown operation command ignored: %s\", buf);",
        "    }",
        "    return 0;",
        "}",
        "if (_text_cb != RT_NULL) {",
        "    _text_cb(buf, (int)pkt_len);",
        "}",
    ])

    # Module 4: 6 pages
    add_page(doc, "模块 4：BLE CSV 分片与帧序列模块", pad([
        "对应源文件：imu_notify_thread.c 中 _send_csv_fragments()。",
        "问题背景：NimBLE 配置的 ATT Preferred MTU 为 247，而完整 CSV 帧约 517 字节。",
        "完整帧格式：[DATA]<timestamp>,<hand>,<seq>,<69 integers>\\n。",
        "单次 BLE Notify 无法稳定承载完整 CSV，因此固件定义 [FRAG] 文本分片协议。",
        "分片格式：[FRAG]<frame_seq>,<hand>,<idx>,<total>,<payload_part>\\n。",
        "字段统计：去掉 [DATA] 前缀后共有 72 个 CSV 字段，包含 timestamp、hand_type、frame_seq、ch0 至 ch9 的 60 个六轴字段以及 ch10 的 9 个九轴字段。",
        "Notify 线程的发送周期为 90 ms，约 11.1 Hz；CSV_BUFFER_SIZE 为 600，用于容纳约 517 字节长帧及余量。",
        "[此处插入分片流程图]",
    ], 50, "分片"))
    add_page(doc, "模块 4：关键常量与长度计算", pad([
        "BLE_NOTIFY_FRAGMENT_SAFE_SIZE = 180。",
        "FRAG_HEADER_MAX_SIZE = 48。",
        "max_payload = 180 - 48 - 1 = 131 字节。",
        "frag_total = ceil(csv_len / max_payload)。",
        "分片发送时先构造 [FRAG] 头部，再复制 payload_part，最后补换行符。",
        "若头部长度或总长度超过安全缓冲区，函数记录告警并返回失败。",
    ], 50, "常量"))
    code_page(doc, "模块 4：分片构造伪代码", [
        "伪代码与源码结构保持一致，强调帧序号、分片序号和 payload 切片关系。",
    ], [
        "max_payload = BLE_NOTIFY_FRAGMENT_SAFE_SIZE - FRAG_HEADER_MAX_SIZE - 1;",
        "frag_total = (csv_len + max_payload - 1) / max_payload;",
        "for (i = 0; i < frag_total; i++) {",
        "    hlen = snprintf(buf, sizeof(buf),",
        "                    \"[FRAG]%u,%s,%u,%u,\",",
        "                    frame_seq, hand_type, i, frag_total);",
        "    part_len = min(max_payload, csv_len - offset);",
        "    memcpy(buf + hlen, csv + offset, part_len);",
        "    buf[hlen + part_len] = '\\n';",
        "    om = ble_hs_mbuf_from_flat(buf, hlen + part_len + 1);",
        "    rc = ble_gattc_notify_custom(conn_handle,",
        "                                  imu_ble_notify_handle(), om);",
        "    offset += part_len;",
        "}",
    ])
    code_page(doc, "模块 4：帧序列临界区保护", [
        "帧序号由 notify 线程和 CMD 控制路径共同访问，因此使用 rt_hw_interrupt_disable/enable 保护。",
    ], [
        "static uint32_t _frame_seq_get(void)",
        "{",
        "    rt_base_t level = rt_hw_interrupt_disable();",
        "    uint32_t seq = _frame_seq;",
        "    rt_hw_interrupt_enable(level);",
        "    return seq;",
        "}",
        "",
        "void imu_notify_reset_frame_seq(void)",
        "{",
        "    rt_base_t level = rt_hw_interrupt_disable();",
        "    _frame_seq = 0;",
        "    rt_hw_interrupt_enable(level);",
        "}",
    ])
    add_page(doc, "模块 4：错误处理与递增语义", pad([
        "_send_csv_fragments 使用 any_sent 判断是否至少有分片进入 notify 调用。",
        "若 mbuf 分配失败，该分片跳过并继续尝试后续分片。",
        "若 ble_gattc_notify_custom 返回 BLE_HS_ENOTCONN，立即中断循环，调用方随后清空连接句柄。",
        "若没有任何分片进入 notify，则返回 1，调用方不递增 _frame_seq。",
        "若任一分片进入 notify，即使后续分片有错误，调用方仍执行 _frame_seq_inc。",
        "因此帧序列语义为：全失败不递增，存在发送尝试则递增。",
    ], 50, "错误"))
    add_page(doc, "模块 4：发送与重组示例", pad([
        "示例原始 CSV：[DATA]102030,left,7,100,101,...,690。",
        "示例分片 1：[FRAG]7,left,0,3,[DATA]102030,left,7,100,101,...。",
        "示例分片 2：[FRAG]7,left,1,3,...中间字段...。",
        "示例分片 3：[FRAG]7,left,2,3,...尾部字段。",
        "上位机以 frame_seq、hand、idx、total 为键缓存分片，idx 从 0 到 total-1 全部到齐后按序拼接 payload_part。",
        "若同一 frame_seq 缺失某个 idx，上位机应丢弃该帧或等待超时后清理缓存。",
        "OpenSpec 约束建议 ROCK 侧按 (hand_type, frame_seq) 重组，FRAG_TIMEOUT_SEC 为 1.0 s，MAX_PENDING_FRAMES 为 128；这些属于上位机消费端约束，不改变固件 Notify 格式。",
        "固件侧没有分片 ACK 或重传机制，采用 BLE Notify 的尽力发送语义，断连时依靠 GAP DISCONNECT 后重新广播恢复链路。",
        "[此处插入分片/重组示例图]",
    ], 50, "示例"))

    # Module 5, 8-12, appendix before TTS/STT
    for title, lines in [
        ("模块 5：运行模式管理模块", [
            "对应源文件：operation_mode.c/.h。",
            "枚举 operation_mode_t 包含 OP_MODE_AUTO 与 OP_MODE_MANUAL。",
            "枚举 operation_state_t 包含 OP_STATE_AUTO_STANDBY、OP_STATE_MANUAL_SLEEP、OP_STATE_RUNNING。",
            "启动默认状态为 AUTO / AUTO_STANDBY，BLE Notify 静默。",
            "Notify 门控由 operation_mode_ble_notify_enabled() 实现，仅 RUNNING 返回 RT_TRUE。",
            "operation_mode.h 明确说明该状态机不影响 TCP upload behavior，只作用于 BLE Notify 输出和 BLE frame_seq 复位。",
            "按钮线程栈大小 1024 字节，优先级 25，轮询 PC6 电平并在稳定低电平后触发本地模式切换。",
            "[此处插入状态机转移图]",
        ]),
        ("模块 5：PC6 按键与消抖", [
            "PC6 按键配置为 PIN_MODE_INPUT_PULLUP，按下时接地为低电平。",
            "按键轮询周期为 20 ms，消抖时间为 60 ms。",
            "源码实现状态机：IDLE -> PRESS_DEBOUNCE -> PRESSED -> RELEASE_DEBOUNCE -> IDLE。",
            "长按只触发一次，避免在 PRESSED 状态重复切换。",
            "本地切换：AUTO -> MANUAL / MANUAL_SLEEP；MANUAL -> AUTO / AUTO_STANDBY。",
            "该按键只影响本端，不要求左右手之间直接通信；双手协调由 ROCK 或外部中心端分别向两端写入 CMD 命令完成。",
        ]),
        ("模块 5：BLE CMD 远程控制", [
            "CMD:RESET_SEQ：调用 imu_notify_reset_frame_seq，状态不变。",
            "CMD:START：重置帧序号并进入 OP_STATE_RUNNING。",
            "CMD:STOP：重置帧序号，若当前为 MANUAL 则进入 MANUAL_SLEEP，否则进入 AUTO_STANDBY。",
            "CMD 控制路径由 Text 特征访问函数先行拦截，避免命令被当作 TTS 文本。",
            "状态切换日志通过 rtdbg.h 分级输出，便于串口调试。",
            "CMD 必须完全匹配字符串；例如 CMD:START 会重置 frame_seq 后进入 RUNNING，普通 SAY: 或自然语言文本不会触发状态切换。",
        ]),
        ("模块 5：状态转移表", [
            "AUTO / AUTO_STANDBY + 本地按键：切换到 MANUAL / MANUAL_SLEEP。",
            "MANUAL / MANUAL_SLEEP + 本地按键：切换到 AUTO / AUTO_STANDBY。",
            "任意模式 + CMD:START：进入 RUNNING。",
            "RUNNING + CMD:STOP：回到当前模式对应的静默状态。",
            "任意状态 + CMD:RESET_SEQ：帧序号归零，模式和状态保持不变。",
        ]),
        ("模块 8：OLED 状态显示模块", [
            "对应源文件：IIC/OLED/*.c、UART/uart_oled_thread.c。",
            "显示器件为 SSD1306/同类 128×64 I2C OLED，源码中 OLED_SCL_PIN 为 PE1，OLED_SDA_PIN 为 PE2。",
            "OLED 库提供 OLED_Init、OLED_Update、OLED_ShowString、OLED_Printf 等接口。",
            "IIC 线程优先级为 20，栈大小为 2048 字节；OLED 位于 TCA9548A 通道 3，若未探测到 TCA9548A，也支持 OLED 直接挂接在 I2C2。",
            "TCA9548A 驱动复用 OLED 的 PE1/PE2 软件 I2C 引脚，并通过 tca9548a_is_present 区分扩展板是否在线。",
            "字库包含 ASCII 8×16、6×8 以及中文点阵数据。",
            "右手 main.c 中封装 oled_show_error_status 用于启动错误显示。",
        ]),
        ("模块 8：显示内容分区", [
            "显示线程可周期刷新连接状态、运行模式、采集状态、电池电量和网络 IP。",
            "右手 STT 识别结果通过 oled_show_stt_result 显示。",
            "UART/OLED 线程负责接收文本并滚动显示到 OLED。",
            "右手工程的 uart_oled 线程用于显示 I2S/STT 文字结果，STT 成功后 voice_assistant 调用 oled_show_stt_result。",
            "OLED 与 I2C2 上的部分传感器共享物理总线，访问时需遵守总线互斥策略。",
            "具体刷新布局、中文字符表和滚动策略详见源代码。",
        ]),
        ("模块 8：错误状态显示 API", [
            "main.c 通过 show_error_on_oled 调用 oled_show_error_status。",
            "当 IIC/OLED 线程、UART/OLED 线程或其他启动环节失败时，可在 OLED 上显示错误状态。",
            "该接口用于现场无调试器场景下的快速故障定位。",
            "错误显示内容应保持简短，适应 128×64 点阵屏显示空间。",
            "后续配图可插入典型启动、连接、低电量和 STT 文本显示截图。",
        ]),
        ("模块 9：电池与电源监测模块", [
            "对应源文件：adc_battery.c/.h。",
            "模块提供 battery_thread_entry、battery_get_voltage、battery_get_percentage 等接口。",
            "ADC 初始化后周期采样电池分压节点，通过平均滤波降低瞬态抖动。",
            "ADC 使用 STM32 HAL 直接操作 ADC1 Channel 12，即 PC2 引脚；ADC 分辨率 12 位，参考电压 3.3 V，分压比为 2。",
            "线程配置为 bat_adc，栈大小 1024 字节，优先级 22，时间片 10；每 1 秒读取一次 ADC 并进行 8 点滑动平均。",
            "adc_to_battery_voltage 将 ADC 原始值换算为电池电压毫伏。",
            "采样结果保存在 g_battery_voltage_mv 和 g_battery_percentage。",
        ]),
        ("模块 9：电量映射与告警", [
            "锂电池电压按 3.0 V 至 4.2 V 区间映射为 0% 至 100%。",
            "源码注释中说明满电 4.1 V 时 ADC 采样点约为 2.05 V，对应 12 位 ADC 原始值约 2546；百分比宏仍按 BATTERY_FULL_VOLTAGE 4200 mV 与 BATTERY_EMPTY_VOLTAGE 3000 mV 线性映射。",
            "低电量阈值用于 OLED 提示、日志输出或后续省电策略触发。",
            "battery_get_voltage 返回毫伏值，battery_get_percentage 返回百分比。",
            "finsh 命令 bat 可显示当前电池电压和百分比。",
            "具体 ADC 通道、分压比例和滤波窗口详见源代码。",
        ]),
        ("模块 10：辅助通信与调试模块", [
            "对应源文件：tcp_client.c/.h、nimble_hci_adapter.c。",
            "nimble_hci_adapter 实现 NimBLE Host 与 CYW43438 HCI 控制器之间的适配。",
            "HCI H4 传输包含命令、ACL、事件等包类型的 TX/RX 状态机。",
            "CYW43438 BT 包提供 cyw43438_bt_init、cyw43438_bt_is_ready、cyw43438_bt_take_rx_ownership、cyw43438_bt_host_send_packet 等接口，供 NimBLE 适配层接管 RX。",
            "UART7 HCI 与 SDIO WiFi 为 CYW43438 上的独立接口；WiFi 继续由 WHD/lwIP 管理，BLE 由 NimBLE Host 管理。",
            "mbuf 内存池由 NimBLE Host 侧使用，notify 路径通过 ble_hs_mbuf_from_flat 创建发送缓冲。",
            "日志系统基于 rtdbg.h，模块使用 DBG_SECTION_NAME 或 DBG_TAG 标识输出来源。",
        ]),
        ("模块 10：TCP 辅助链路", [
            "tcp_client.c 用于辅助上传 IMU、电池或翻译文本等数据，具体启用方式以工程配置为准。",
            "TCP 默认服务器 IP 为 192.168.221.92，端口 8266；发送缓冲区 4096 字节，接收缓冲区 1024 字节，发送间隔 100 ms。",
            "TCP 客户端线程名 tcp_cli，优先级 22，栈大小 4096 字节；支持 MSH 命令 tcp_client_start/tcp_client_stop 及别名 tcp_start/tcp_stop。",
            "左手端文本链路中 _ble_text_recv_handler 调用 tcp_set_translated_text 保存规范化文本。",
            "左手 tcp_client 还提供 translated_text 暂存和发送成功后清理逻辑，用于把 BLE/TTS 收到的翻译文本同步到 TCP JSON 辅助链路。",
            "TCP 链路不替代 BLE 申报核心功能，主要用于调试、扩展或与外部工具协同。",
            "网络相关功能依赖 CYW43438 WiFi 与 WHD 驱动。",
            "如提交鉴别材料时不包含外部服务器，应说明该链路为辅助调试功能。",
        ]),
        ("模块 10：finsh 与日志", [
            "RT-Thread finsh shell 可用于导出命令、查看状态和触发测试。",
            "adc_battery.c 中 MSH_CMD_EXPORT(bat, ...) 导出电池查看命令。",
            "VTX316 模块也提供命令行测试入口，便于串口验证播报。",
            "VTX316 命令为 vtx316_say <text>；TCP 命令为 tcp_start [ip] [port] 与 tcp_stop；WiFi 可通过 RT-Thread WLAN MSH 命令连接。",
            "日志级别按 DBG_INFO、DBG_DEBUG 等宏控制。",
            "调试日志不影响核心业务协议，发布版本可按配置降低输出。",
        ]),
        ("模块 11：硬件接口与引脚分配表", [
            "主控：STM32H7R7，Cortex-M7 内核。",
            "无线：CYW43438 WiFi/BT Combo 模组。",
            "I2C2：PE1-SCL、PE2-SDA，连接 TCA9548A、OLED 和部分 IMU。",
            "I2C1：PB8-SCL、PB9-SDA，连接 TCA9548A #1 及 8 路 MPU6050。",
            "UART1：PF13-TX、PF12-RX，左手用于 VTX316，右手用于外部串口文本输入/OLED 显示。",
            "UART4：工程注释中作为控制台/MSH Shell，常用 115200 baud。",
            "UART7：CYW43438 蓝牙 HCI H4 传输，默认 115200 baud。",
            "SAI/INMP441：源码注释显示 SCK 为 PA2、WS 为 PC0，SD 存在 PE3 与 PE7 两种接线说明，需以实际硬件焊接为准。",
            "GPIO：PC6 按键，上拉输入，按下接地。",
            "ADC：用于锂电池电压分压采样。",
            "[此处插入 CYW43438 连接框图]",
        ]),
        ("模块 11：接口表", [
            "I2C1/I2C2：用于 MPU6050、ICM-20948、TCA9548A 与 OLED。",
            "SAI：用于右手 INMP441 数字麦克风音频输入。",
            "UART7：用于 CYW43438 蓝牙 HCI 或相关串口链路，具体实现详见源代码。",
            "UART：用于左手 VTX316 语音合成模块。",
            "SDIO：用于 CYW43438 WiFi 数据通路，具体实现详见 BSP 与 WHD 驱动。",
        ]),
        ("模块 12：术语与参考文档", [
            "BLE：Bluetooth Low Energy，低功耗蓝牙。",
            "GATT：Generic Attribute Profile，BLE 属性访问框架。",
            "ATT：Attribute Protocol，GATT 底层属性协议。",
            "MTU：Maximum Transmission Unit，单次 ATT PDU 可承载的最大长度。",
            "NimBLE：Apache Mynewt 项目的 BLE 协议栈，本工程使用 V1.0.0。",
        ]),
        ("模块 12：组件与许可证", [
            "RT-Thread：嵌入式实时操作系统，Apache 2.0。",
            "STM32 HAL：STMicroelectronics HAL 驱动，BSD 3-Clause。",
            "CMSIS：Arm Cortex 微控制器软件接口标准，Apache 2.0。",
            "NimBLE：BLE Host 协议栈，Apache 2.0。",
            "SCons、Keil MDK、ARM Compiler 等工具用于构建，不作为本软件运行时功能申报。",
        ]),
        ("附录 A：测试与验证", [
            "测试 1：左右手设备上电启动，串口日志显示 RT-Thread 与主线程启动。",
            "测试 2：CYW43438 HCI Reset 完成后，BLE 广播名可被扫描到。",
            "测试 3：上位机连接 ART-Pi2-IMU-L 并订阅 Notify 特征。",
            "测试 4：上位机连接 ART-Pi2-IMU-R 并订阅 Notify 特征。",
            "测试 5：CMD:START 后，Notify 开始输出 [FRAG] 分片。",
            "测试 6：CMD:STOP 后，Notify 停止发送。",
        ]),
        ("附录 A：BLE 连接稳定性测试", [
            "测试 7：反复断开和重连 BLE，设备应自动重新广播。",
            "测试 8：订阅和取消订阅 Notify，_subscribed 状态应正确变化。",
            "测试 9：MTU 更新事件应记录协商结果，不影响分片协议。",
            "测试 10：连接句柄失效时，BLE_HS_ENOTCONN 应触发发送中断并清理句柄。",
            "测试 11：左右手同时运行时，设备名应能区分 L/R。",
        ]),
        ("附录 A：分片重组正确性测试", [
            "测试 12：采集一帧 69 整数字段 CSV，验证分片 total 与 payload 长度计算正确。",
            "测试 13：按 frame_seq、hand、idx、total 重组后，应恢复原始 [DATA] 帧。",
            "测试 14：人为丢弃一个 idx，上位机应超时丢弃该帧。",
            "测试 15：发送 CMD:RESET_SEQ 后，下一帧 frame_seq 应从 0 重新开始。",
            "测试 16：IMU 某通道异常时，CSV 字段数保持不变，对应位置为 0。",
        ]),
        ("附录 B：配置与部署", [
            "构建环境：Windows 11、Keil MDK-ARM μVision5、ARM Compiler V6.22、SCons、RT-Thread ENV。",
            "构建步骤：进入 art_pi2_left 或 art_pi2_right 工程目录，使用 ENV/SCons 生成或更新 Keil 工程。",
            "rtconfig.py 中 CROSS_TOOL 默认值为 gcc，但可通过 RTT_CC=keil 切换到 armclang；当使用 Keil MDK 时 PLATFORM 为 armclang，EXEC_PATH 指向 Keil_v5。",
            "applications/SConscript 默认编译该目录下全部 .c 文件；IIC/SConscript 编译 IIC 目录和 OLED 子目录；vtx316/SConscript 依赖 BSP_USING_UART1。",
            "打开 Keil 工程后选择目标配置，执行编译并生成固件镜像。",
            "烧录方式：通过 ST-Link 或 J-Link 连接 ART-Pi2 调试接口进行下载。",
            "串口日志：连接调试串口，按工程配置波特率查看 rt_kprintf 与 rtdbg 输出。",
        ]),
        ("附录 B：运行配置", [
            "BLE：确认 CYW43438 模组供电、时钟和 HCI 通路正常。",
            "WiFi：确认 WHD 固件、CLM、NVRAM 资源分区或外部存储文件可访问，并确认 WLAN 设备名 wlan0 存在。",
            "IMU：确认 TCA9548A 地址、I2C 总线和 11 路传感器连接正确。",
            "TTS：左手端确认 VTX316 串口接线和电源正常。",
            "STT：右手端确认 INMP441 SAI 接线、WiFi 配置和讯飞代理服务地址。",
            "云服务密钥：APPID、APIKey、APISecret 等应在提交材料中脱敏。",
        ]),
    ]:
        add_page(doc, title, pad(lines, 50, title.split("：")[0]))

    code_page(doc, "补充代码摘录：运行模式与本地按键", [
        "以下摘录对应 operation_mode.c/.h，展示 PC6 按键、消抖参数和 BLE Notify 门控的核心实现。",
    ], [
        "#define OP_MODE_BUTTON_PIN          GET_PIN(C, 6)",
        "#define OP_MODE_BUTTON_POLL_MS      20",
        "#define OP_MODE_BUTTON_DEBOUNCE_MS  60",
        "",
        "rt_bool_t operation_mode_ble_notify_enabled(void)",
        "{",
        "    return (_state == OP_STATE_RUNNING) ? RT_TRUE : RT_FALSE;",
        "}",
        "",
        "if (rt_strcmp(cmd, \"CMD:START\") == 0) {",
        "    imu_notify_reset_frame_seq();",
        "    _state = OP_STATE_RUNNING;",
        "    return RT_TRUE;",
        "}",
        "if (rt_strcmp(cmd, \"CMD:STOP\") == 0) {",
        "    imu_notify_reset_frame_seq();",
        "    _state = (_mode == OP_MODE_MANUAL) ?",
        "             OP_STATE_MANUAL_SLEEP : OP_STATE_AUTO_STANDBY;",
        "    return RT_TRUE;",
        "}",
    ])

    code_page(doc, "补充代码摘录：OLED 与 TCA9548A", [
        "以下摘录对应 IIC/tca9548a.h、IIC/iic_thread.c、OLED/OLED.h，展示 I2C2 OLED/TCA 通道配置。",
    ], [
        "#define TCA9548A_ADDR       0x70",
        "#define TCA9548A_MAX_CHANNEL    8",
        "#define OLED_TCA9548A_CHANNEL   3",
        "#define OLED_SCL_PIN    GET_PIN(E, 1)",
        "#define OLED_SDA_PIN    GET_PIN(E, 2)",
        "",
        "OLED_I2C_Init();",
        "tca9548a_init();",
        "if (tca9548a_is_present()) {",
        "    tca9548a_select_channel(OLED_TCA9548A_CHANNEL);",
        "}",
        "OLED_Init();",
        "OLED_Update();",
        "oled_show_system_status(\"OLED READY\");",
    ])

    code_page(doc, "补充代码摘录：电池 ADC 与 TCP 参数", [
        "以下摘录对应 adc_battery.h、tcp_client.h，展示电池采样和辅助 TCP 链路的关键常量。",
    ], [
        "#define BATTERY_THREAD_STACK_SIZE   1024",
        "#define BATTERY_THREAD_PRIORITY     22",
        "#define BATTERY_FULL_VOLTAGE        4200",
        "#define BATTERY_EMPTY_VOLTAGE       3000",
        "#define BATTERY_DIVIDER_RATIO       2",
        "#define ADC_REF_VOLTAGE             3300",
        "#define ADC_RESOLUTION              4096",
        "",
        "#define TCP_SERVER_IP       \"192.168.221.92\"",
        "#define TCP_SERVER_PORT     8266",
        "#define TCP_THREAD_STACK_SIZE   4096",
        "#define TCP_SEND_BUF_SIZE       4096",
        "#define TCP_RECV_BUF_SIZE       1024",
        "#define TCP_SEND_INTERVAL       100",
    ])

    # Module 6: 3 pages near the end
    add_page(doc, "模块 6：左手语音播报模块（TTS）", pad([
        "对应源文件：vtx316/vtx316.c/.h、ble_app.c 的 _ble_text_recv_handler。",
        "VTX316 模块通过 UART 与左手 ART-Pi2 通信，用于将文本转换为语音播报。",
        "vtx316.h 定义串口设备名为 uart1，波特率 115200，接线为 USART1 PF13-TX、PF12-RX。",
        "协议帧格式为 0xFD、长度高字节、长度低字节、命令 0x01、编码 0x05、UTF-8 文本；播报完成回复为 0x41 0x4F，即 ASCII \"AO\"。",
        "头文件暴露 vtx316_speak 与 vtx316_speak_wait 两个接口。",
        "vtx316_speak 为非阻塞接口，模块忙时可能跳过；vtx316_speak_wait 阻塞等待上一次播报结束后再发送。",
        "使用场景：把健听人语音识别结果或上位机翻译文本朗读给听障用户。",
        "[此处插入 TTS 链路图]",
    ], 50, "TTS"))
    add_page(doc, "模块 6：BLE 到 TTS 链路", pad([
        "链路顺序：BLE Text 写入 -> _text_chr_access -> CMD: 拦截 -> _text_cb -> _ble_text_recv_handler -> 文本规范化 -> vtx316_speak_wait。",
        "Text 特征最大接受 256 字节 UTF-8 文本。",
        "若文本以 CMD: 开头，直接进入 operation_mode_handle_cmd，不进入 TTS。",
        "普通文本由 _ble_text_recv_handler 处理，并通过 tcp_set_translated_text 保存辅助链路文本。",
        "播报调用采用同步等待方式，避免连续文本覆盖或串口命令交叠。",
        "vtx316_speak_wait 在模块 busy 时等待 VTX316_EVT_DONE，超时时间为 30000 ms；若超时则强制清 busy 后发送新帧。",
        "VTX316 接收回调在中断上下文检测 0x41 0x4F 序列，收到完整完成标志后释放 rt_event。",
    ], 50, "TTS链路"))
    code_page(doc, "模块 6：文本规范化片段", [
        "规范化逻辑支持 SAY: 前缀、大小写 SAY、半角冒号、全角冒号和首尾空白处理。",
    ], [
        "if (text_len >= 3 &&",
        "    (start[0] == 'S' || start[0] == 's') &&",
        "    (start[1] == 'A' || start[1] == 'a') &&",
        "    (start[2] == 'Y' || start[2] == 'y')) {",
        "    start += 3;",
        "    text_len -= 3;",
        "    while (text_len > 0) {",
        "        if (*start == ':' || *start == ' ' || *start == '\\t') {",
        "            start++; text_len--; continue;",
        "        }",
        "        break;",
        "    }",
        "}",
        "vtx316_speak_wait(normalized_text);",
    ])

    # Module 7: 5 pages + confirmation page = 60 body pages
    for title, lines in [
        ("模块 7：右手音频采集与 AI 云识别模块（STT）", [
            "对应源文件：drv_sai_inmp441.c、audio_capture_inmp441.c、audio_process.c、voice_assistant.c、ai_cloud_service.c、web_client.c。",
        "INMP441 为数字麦克风，输出 24-bit I2S/PCM 数据。",
        "drv_sai_inmp441.h 定义 INMP441_SAMPLE_RATE = 16000、INMP441_BIT_WIDTH = 24、INMP441_CHANNEL_NUM = 1。",
        "音频缓冲配置为 SAI_DMA_BUFFER_SIZE = 1024、AUDIO_BUFFER_COUNT = 4、AUDIO_FRAME_SIZE = 512。",
        "工程目标格式为 16 kHz、16-bit、mono、little-endian PCM，用于语音识别接口。",
            "drv_sai_inmp441.h 中定义 INMP441_SAMPLE_RATE = 16000、INMP441_BIT_WIDTH = 24、INMP441_CHANNEL_NUM = 1。",
            "[此处插入 STT 数据流图]",
        ]),
        ("模块 7：SAI 与 DMA 采集", [
            "源码注释中 SAI2 SCK 使用 PA2，WS 使用 PC0，SD 接线在头文件与实现注释中存在 PE3/PE7 差异，需按实际硬件确认。",
            "SAI_DMA_BUFFER_SIZE 定义为 1024 samples。",
            "驱动层通过 inmp441_init、inmp441_start、inmp441_read_frame 和 inmp441_stop 提供采集接口。",
        "audio_capture_inmp441.c 将 INMP441 24-bit 有效数据转换为 16-bit PCM。",
        "inmp441_device_t 包含 buffer_sem、lock、frames[4]、write_idx、read_idx、frame_count、total_frames、overrun_count、dma_errors 等状态字段。",
        "inmp441_read_frame 为阻塞式读取接口，调用方需释放 frame->buffer。",
        "DMA 双缓冲和底层 SAI 中断细节具体实现详见 drv_sai_inmp441.c。",
        ]),
        ("模块 7：音频处理与 VAD", [
            "audio_process_init 分配最大录音时长对应的 recording buffer。",
        "audio_process_start 每次创建 audio_proc 线程，并重置录音状态、VAD hangover 计数和降噪状态。",
        "audio_process 线程栈大小 4096 字节，优先级 10；VAD_THRESHOLD = 20000，VAD_HANGOVER_FRAMES = 31，VAD_MIN_SPEECH_FRAMES = 10。",
        "最大录音时长 AUDIO_PROCESS_MAX_RECORD_SEC = 2，源码注释说明每秒 16000 samples × 4 bytes = 64 KB，2 秒约 128 KB。",
        "audio_process_thread_entry 循环读取 audio_frame_t，执行 noise reduction、energy 计算和 vad_detect_speech。",
            "状态从 AUDIO_STATE_IDLE 进入 AUDIO_STATE_RECORDING 后缓存语音帧，端点结束后通过回调交给 voice_assistant。",
            "speech_data_handler 将 32-bit 采样转换为 16-bit PCM，并释放 speech_ready_sem 通知主任务。",
        ]),
        ("模块 7：语音助手任务编排", [
            "voice_assistant_init 依次初始化 audio_capture、audio_process 和 ai_cloud_service。",
            "voice_assistant_start 创建 voice_asst 主线程。",
        "主线程启动 audio_capture_start 与 audio_process_start，进入 VOICE_ASSISTANT_LISTENING。",
        "voice_assistant 内部维护 VOICE_ASSISTANT_IDLE、LISTENING、PROCESSING 等状态，使用 trigger_sem 与 speech_ready_sem 协调录音完成事件。",
        "若录音数据小于 3200 字节，即约 100 ms 的 16 kHz 16-bit 单声道 PCM，主任务会跳过该段，避免过短音频触发 STT。",
        "VAD 录音完成后，线程取出 pcm16_buffer，重启后台采集与处理，再调用 ai_cloud_service_speech_to_text。",
            "识别成功后通过 oled_show_stt_result 在 OLED 上显示文字。",
        ]),
        ("模块 7：讯飞开放平台 API 集成", [
            "voice_assistant_config.h 以 AI_SERVICE_PROVIDER 选择云服务提供商，其中 1 表示 iFlytek。",
        "源码当前采用 XFYUN_STT_URL 指向 PC 代理服务 tools/xfyun_proxy.py。",
        "voice_assistant_config.h 中 XFYUN_STT_URL 当前为 http://192.168.221.92:8080/stt，VOICE_STT_ENABLE = 1，VOICE_TTS_ENABLE = 0，VOICE_FULL_DUPLEX_ENABLE = 0。",
        "该配置文件含 APPID、APIKey、APISecret 示例/实值，正式提交文档与源码鉴别材料时应做脱敏处理。",
        "ai_cloud_service_speech_to_text 在讯飞代理模式下以 application/octet-stream 发送原始 PCM，避免在 MCU 端构造大 Base64 JSON。",
            "代理服务负责与讯飞开放平台接口交互，包括 APPID/APIKey/APISecret、HMAC 签名、实时语音识别 IAT 数据流和 JSON 响应解析。",
            "MCU 端 parse_stt_response 期望代理返回 {\"result\":[\"识别文本\"]} 格式，并提取 text_result。",
        ]),
    ]:
        add_page(doc, title, pad(lines, 50, "STT"))

    code_page(doc, "补充代码摘录：INMP441 与 VAD 参数", [
        "以下摘录对应 drv_sai_inmp441.h、audio_process.h，展示音频采集、帧缓冲和 VAD 端点检测的关键常量。",
    ], [
        "#define INMP441_SAMPLE_RATE         16000",
        "#define INMP441_BIT_WIDTH           24",
        "#define INMP441_CHANNEL_NUM         1",
        "#define SAI_DMA_BUFFER_SIZE         1024",
        "#define AUDIO_BUFFER_COUNT          4",
        "#define AUDIO_FRAME_SIZE            512",
        "",
        "#define AUDIO_PROCESS_STACK_SIZE        4096",
        "#define AUDIO_PROCESS_PRIORITY          10",
        "#define VAD_THRESHOLD                   20000",
        "#define VAD_HANGOVER_FRAMES             31",
        "#define VAD_MIN_SPEECH_FRAMES           10",
        "#define AUDIO_PROCESS_MAX_RECORD_SEC    2",
    ])

    add_page(doc, "待确认事项", pad([
        "[待确认：右手 INMP441 实际 SD 引脚以硬件焊接为准；源码注释中存在 PE3 与 PE7 两种说明。]",
        "[待确认：讯飞开放平台 WebSocket/HMAC/IAT 细节当前主要位于 PC 代理 tools/xfyun_proxy.py，若需作为固件直接实现申报，应补充 MCU 端源码。]",
        "[待确认：voice_assistant_config.h 当前含讯飞 APPID/APIKey/APISecret 等敏感字段，提交软件著作权鉴别材料前应确认是否需要替换为占位符。]",
        "[待确认：BLE 配置文档记录的 MAC 地址来自 CYW43438 芯片固件，若批量更换开发板，最终提交版本应更新实测 MAC 或删除固定 MAC 描述。]",
        "[待确认：VTX316 串口号、波特率和完整帧格式请以最终硬件接线与 vtx316.c 配置为准。]",
        "[待确认：ADC 电池分压比例、低电量阈值和最终电池容量参数请以实测硬件版本为准。]",
        "[待确认：ROCK 5B 上位机识别程序不在本次申报范围内，提交材料时应避免混入其源码页。]",
    ], 50, "待确认"))

    doc.save(OUT)


if __name__ == "__main__":
    main()
