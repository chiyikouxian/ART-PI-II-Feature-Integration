from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = "document_identification.docx"
SOFTWARE = "可穿戴双向手语-语音交互固件系统"
VERSION = "V1.0"
HEADER = f"{SOFTWARE} {VERSION}"


def font(run, name="宋体", size=10.5, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def field(paragraph, instr):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = instr
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, code, sep, text, end])
    font(run, size=9)


def setup(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(1)
    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        st = doc.styles[name]
        st.font.name = "宋体"
        st.font.size = Pt(size)
        st.font.bold = True
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        st.paragraph_format.space_before = Pt(5)
        st.paragraph_format.space_after = Pt(2)
    if "CodeBlock" not in [s.name for s in doc.styles]:
        code = doc.styles.add_style("CodeBlock", 1)
    else:
        code = doc.styles["CodeBlock"]
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.space_after = Pt(0)


def section(section, footer=False, restart=False):
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)
    hp = section.header.paragraphs[0]
    hp.text = HEADER
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs:
        font(r, size=9)
    if restart:
        pg = section._sectPr.find(qn("w:pgNumType"))
        if pg is None:
            pg = OxmlElement("w:pgNumType")
            section._sectPr.append(pg)
        pg.set(qn("w:start"), "1")
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = ""
    if footer:
        fp.add_run("第 ")
        field(fp, "PAGE")
        fp.add_run(" 页 / 共 ")
        field(fp, "SECTIONPAGES")
        fp.add_run(" 页")
        for r in fp.runs:
            font(r, size=9)


def p(doc, text="", style="Normal", align=None, bold=False):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    font(run, "宋体", 10.5, bold)
    return para


def lines(doc, arr):
    for line in arr:
        p(doc, line)


def code(doc, arr):
    p(doc, "参考代码：", bold=True)
    p(doc, "```c", "CodeBlock")
    for line in arr:
        para = doc.add_paragraph(style="CodeBlock")
        run = para.add_run(line)
        font(run, "Consolas", 9)
    p(doc, "```", "CodeBlock")


def table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(val)
            font(r, bold=(i == 0))


def cover(doc):
    for _ in range(5):
        p(doc)
    p(doc, SOFTWARE, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True).runs[0].font.size = Pt(22)
    p(doc, VERSION, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True).runs[0].font.size = Pt(18)
    p(doc, "软件说明书", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True).runs[0].font.size = Pt(18)
    for _ in range(7):
        p(doc)
    for line in ["软件名称：可穿戴双向手语-语音交互固件系统", "版本号：V1.0", "文档类型：软件说明书", "编写日期：【待填写】", "著作权人：【待填写】"]:
        p(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def front(doc):
    doc.add_heading("修订历史", 1)
    table(doc, [["版本", "日期", "修订内容", "编写", "审核"], ["V1.0", "【待填写】", "形成软件著作权申报用软件说明书。", "【待填写】", "【待填写】"]])
    lines(doc, [
        "文档用途：作为中国版权保护中心软件著作权登记的文档鉴别材料。",
        "说明范围：ART-Pi2 左手端与右手端嵌入式固件，包含 IMU、BLE、TTS、STT、OLED、电池监测与辅助 TCP 调试链路；ROCK 5B 上位机识别程序不在本次申报范围内。",
        "敏感信息提示：WiFi 口令、讯飞 APPID/APIKey/APISecret 等配置在提交版本中应脱敏。",
    ])
    doc.add_page_break()
    doc.add_heading("目录", 1)
    for line in [
        "第 1 章 概述",
        "第 2 章 系统架构",
        "第 3 章 功能模块详述：模块 1 至模块 12",
        "附录 A 测试与验证",
        "附录 B 配置与部署",
        "待确认事项",
    ]:
        p(doc, line)
    doc.add_page_break()


def module(doc, title, body, refs=None):
    doc.add_heading(title, 2)
    lines(doc, body)
    if refs:
        code(doc, refs)


def main():
    doc = Document()
    setup(doc)
    section(doc.sections[0], footer=False)
    cover(doc)
    front(doc)

    body = doc.add_section(WD_SECTION_START.NEW_PAGE)
    body.header.is_linked_to_previous = False
    body.footer.is_linked_to_previous = False
    section(body, footer=True, restart=True)

    doc.add_heading("第 1 章 概述", 1)
    lines(doc, [
        f"软件名称：{SOFTWARE}；版本号：{VERSION}。",
        "软件类型：嵌入式固件系统。运行平台为 ART-Pi2 开发板，主控为 STM32H7R7 系列 MCU，RTOS 为 RT-Thread V5.x。",
        "左手 BLE 名称为 ART-Pi2-IMU-L，右手 BLE 名称为 ART-Pi2-IMU-R。BLE 配置文档记录左手 MAC 为 C0:4E:51:05:34:33，右手 MAC 为 C0:31:51:05:34:33；MAC 由 CYW43438 固件提供，应用层不直接写入。",
        "开发目的：为听障人士与健听人群提供双向交流辅助。手语方向由双手 IMU 数据经 BLE 上行给上位机识别，语音方向由右手 INMP441/STT 链路将现场语音转写为文字，左手 VTX316 可把下行文本播报为语音。",
        "开发与构建环境：Windows 11、Keil MDK-ARM μVision5、ARM Compiler V6.22、SCons、RT-Thread ENV，主要语言为 C。",
        "源码统计：自写源代码约 18,702 行，去空行注释约 11,680 行，自写 .c/.h 文件约 85 个；左手约 35 个文件、右手约 50 个文件。",
    ])

    doc.add_heading("第 2 章 系统架构", 1)
    lines(doc, [
        "硬件架构：左右手均使用 ART-Pi2 + CYW43438 WiFi/BT Combo；共有 10 路 MPU6050、1 路 ICM-20948、SSD1306/OLED、PC6 按键、电池 ADC。左手独有 VTX316 TTS，右手独有 INMP441 数字麦克风。",
        "软件分层：应用层包含 main、BLE 应用、IMU Notify、operation_mode、OLED、TTS/STT；协议层包含 NimBLE GATT、[DATA]/[FRAG]、CMD、HTTP 代理；驱动层包含 CYW43438 HCI、I2C/TCA9548A、MPU6050、ICM-20948、SAI/INMP441、UART、ADC。",
        "WiFi 与蓝牙分工：CYW43438 WiFi 走 SDIO/WHD/lwIP，BLE 走 UART7 HCI H4/NimBLE Host。rtconfig.h 启用 WLAN、SAL、netdev、lwIP 2.1.2、POSIX socket、FAL、DFS、finsh/MSH。",
        "主启动关系：左手 main 创建 iic_drv、vtx316、mpu6050、bat_adc，并在 WiFi 成功后启动 TCP；右手 main 创建 iic_drv、uart_oled、mpu6050、bat_adc、autostart，autostart 负责 WiFi、TCP、voice_assistant 编排。",
        "[此处插入系统架构图]",
    ])

    doc.add_heading("第 3 章 功能模块详述", 1)

    module(doc, "模块 1：系统启动与初始化模块", [
        "对应源文件：main.c、ble_app.c、cyw43438_bt、nimble_hci_adapter。",
        "RT-Thread 初始化使用 BOARD、DEVICE、COMPONENT、ENV、APP 阶段。vtor_config 通过 INIT_BOARD_EXPORT 将中断向量表重定位到 XSPI2_BASE。BLE 应用初始化位于 INIT_ENV_EXPORT(_ble_app_init)。",
        "_ble_app_init 先初始化 operation_mode，再等待 CYW43438 BT ready，随后安装 HCI 适配器、配置 NimBLE Host 回调、注册 GAP/GATT 和自定义 IMU 服务、启动 Host 线程与 Notify 线程。",
        "左手端还注册 BLE Text 回调到 _ble_text_recv_handler，用于 VTX316 播报；左右手分别使用 g_left_ble_init_ok/g_right_ble_init_ok 表示 BLE 初始化结果。",
        "[此处插入启动时序图]",
    ], [
        "static int _ble_app_init(void)",
        "{",
        "    operation_mode_init();",
        "    while (!cyw43438_bt_is_ready() && waited < 100) {",
        "        rt_thread_mdelay(50); waited++;",
        "    }",
        "    if (!cyw43438_bt_is_ready()) return -RT_ETIMEOUT;",
        "    ret = nimble_hci_adapter_init();",
        "    ble_hs_cfg.sync_cb = _on_sync;",
        "    ble_hs_cfg.reset_cb = _on_reset;",
        "    ble_svc_gap_init(); ble_svc_gatt_init();",
        "    ble_svc_gap_device_name_set(DEVICE_NAME);",
        "    ret = imu_ble_service_init();",
        "    ble_hs_thread_startup();",
        "    return imu_notify_thread_start();",
        "}",
    ])

    module(doc, "模块 2：IMU 数据采集模块", [
        "对应源文件：mpu6050/*.c/.h、IIC/*.c/.h。",
        "I2C1 PB8/PB9 连接 TCA9548A #1，ch0-ch7 为 8 路 MPU6050；I2C2 PE1/PE2 连接 TCA9548A #2，ch0-ch1 为 2 路 MPU6050，ch2 为 ICM-20948，ch3 为 OLED。",
        "mpu_channel_data_t 包含 ax/ay/az/gx/gy/gz/mx/my/mz、valid、has_mag。BLE 发送使用 mpu_get_channel_raw_data 保持原始整数语义，异常通道按固定字段数补 0。",
        "ICM-20948 使用 Bank 寄存器访问和 bypass 模式访问 AK09916 磁力计；AK09916 通过 WIA2、CNTL2、CNTL3 等寄存器完成身份检测和测量模式配置。",
        "[此处插入 IMU 阵列布局图]",
    ], [
        "#define MPU_TOTAL_CHANNELS   (MPU6050_I2C1_COUNT + MPU6050_I2C2_COUNT)",
        "typedef struct {",
        "    short ax, ay, az;",
        "    short gx, gy, gz;",
        "    short mx, my, mz;",
        "    rt_bool_t valid;",
        "    rt_bool_t has_mag;",
        "} mpu_channel_data_t;",
        "",
        "if (mpu_get_channel_raw_data(i, &d) == RT_EOK && d.valid) {",
        "    append_sensor_values(&d);",
        "} else {",
        "    append_zero_values_for_channel(i);",
        "}",
    ])

    module(doc, "模块 3：BLE 无线通信模块", [
        "对应源文件：imu_ble_service.c/.h、ble_app.c、syscfg.h。",
        "NimBLE 作为 Host-only 运行，CYW43438 内部 BT 控制器通过 UART7 HCI H4 连接。syscfg.h 中 ATT Preferred MTU 为 247，最大连接数为 1。",
        "自定义服务 UUID 为 A74D0001-B4E7-4C5F-9D2A-F163E80ACB00；Notify、Channel、Text 特征分别使用 A74D0002、A74D0003、A74D0004。",
        "广播模式为 connectable undirected，间隔 200-500 ms。GAP 事件处理 CONNECT、DISCONNECT、SUBSCRIBE、MTU、CONN_UPDATE，断连后自动重新广播。",
        "[此处插入 GATT 服务表] [此处插入 BLE 连接流程图]",
    ], [
        "static const struct ble_gatt_svc_def _gatt_svcs[] = {",
        "  { .type = BLE_GATT_SVC_TYPE_PRIMARY,",
        "    .uuid = &_svc_uuid.u,",
        "    .characteristics = (struct ble_gatt_chr_def[]) {",
        "      { .uuid = &_notify_chr_uuid.u,",
        "        .flags = BLE_GATT_CHR_F_READ | BLE_GATT_CHR_F_NOTIFY,",
        "        .val_handle = &_notify_handle },",
        "      { .uuid = &_channel_chr_uuid.u,",
        "        .flags = BLE_GATT_CHR_F_READ | BLE_GATT_CHR_F_WRITE },",
        "      { .uuid = &_text_chr_uuid.u,",
        "        .flags = BLE_GATT_CHR_F_WRITE },",
        "      { 0 } } },",
        "  { 0 }",
        "};",
    ])

    module(doc, "模块 4：BLE CSV 分片与帧序列模块", [
        "对应源文件：imu_notify_thread.c 中 _send_csv_fragments()。",
        "完整帧格式为 [DATA]<timestamp>,<hand>,<seq>,<69 integers>。去掉 [DATA] 前缀后共 72 个 CSV 字段。CSV_BUFFER_SIZE 为 600，Notify 周期为 90 ms。",
        "由于 ATT MTU 247 小于完整 CSV 长度，固件定义 [FRAG]<frame_seq>,<hand>,<idx>,<total>,<payload_part> 分片格式。BLE_NOTIFY_FRAGMENT_SAFE_SIZE 为 180，FRAG_HEADER_MAX_SIZE 为 48，max_payload 为 131。",
        "frame_seq 使用 rt_hw_interrupt_disable/enable 保护；若没有任何分片进入 notify 调用，则不递增；若存在发送尝试，则帧序列递增一次。BLE_HS_ENOTCONN 触发发送中断和连接句柄清理。",
        "[此处插入分片流程图] [此处插入分片/重组示例图]",
    ], [
        "#define BLE_NOTIFY_FRAGMENT_SAFE_SIZE 180",
        "#define FRAG_HEADER_MAX_SIZE          48",
        "int max_payload = BLE_NOTIFY_FRAGMENT_SAFE_SIZE - FRAG_HEADER_MAX_SIZE - 1;",
        "uint16_t frag_total = (csv_len + max_payload - 1) / max_payload;",
        "hlen = rt_snprintf(frag_buf, sizeof(frag_buf),",
        "                   \"[FRAG]%u,%s,%u,%u,\",",
        "                   frame_seq, hand_type, i, frag_total);",
        "memcpy(frag_buf + hlen, csv + offset, part_len);",
        "frag_buf[hlen + part_len] = '\\n';",
        "om = ble_hs_mbuf_from_flat(frag_buf, hlen + part_len + 1);",
        "rc = ble_gattc_notify_custom(conn_handle, imu_ble_notify_handle(), om);",
    ])

    module(doc, "模块 5：运行模式管理模块", [
        "对应源文件：operation_mode.c/.h。",
        "状态机包含 OP_MODE_AUTO、OP_MODE_MANUAL，以及 OP_STATE_AUTO_STANDBY、OP_STATE_MANUAL_SLEEP、OP_STATE_RUNNING。启动默认 AUTO/AUTO_STANDBY。",
        "PC6 按键采用上拉输入、按下接地；轮询周期 20 ms，消抖 60 ms，长按只触发一次。BLE Notify 门控只在 RUNNING 状态返回真。",
        "Text 特征中的 CMD:RESET_SEQ、CMD:START、CMD:STOP 由 operation_mode_handle_cmd 处理，且该状态机只影响 BLE Notify 和 frame_seq，不影响 TCP 上传。",
        "[此处插入状态机转移图]",
    ], [
        "#define OP_MODE_BUTTON_PIN          GET_PIN(C, 6)",
        "#define OP_MODE_BUTTON_POLL_MS      20",
        "#define OP_MODE_BUTTON_DEBOUNCE_MS  60",
        "rt_bool_t operation_mode_ble_notify_enabled(void)",
        "{",
        "    return (_state == OP_STATE_RUNNING) ? RT_TRUE : RT_FALSE;",
        "}",
        "if (rt_strcmp(cmd, \"CMD:RESET_SEQ\") == 0) imu_notify_reset_frame_seq();",
        "if (rt_strcmp(cmd, \"CMD:START\") == 0) {",
        "    imu_notify_reset_frame_seq(); _state = OP_STATE_RUNNING;",
        "}",
        "if (rt_strcmp(cmd, \"CMD:STOP\") == 0) {",
        "    imu_notify_reset_frame_seq();",
        "    _state = (_mode == OP_MODE_MANUAL) ? OP_STATE_MANUAL_SLEEP : OP_STATE_AUTO_STANDBY;",
        "}",
    ])

    module(doc, "模块 6：左手语音播报模块（TTS）", [
        "对应源文件：vtx316/vtx316.c/.h、ble_app.c 的 _ble_text_recv_handler。",
        "VTX316 通过 UART1 与左手 ART-Pi2 通信，设备名 uart1，波特率 115200，接线为 PF13-TX、PF12-RX。",
        "协议帧为 0xFD、长度高字节、长度低字节、命令 0x01、编码 0x05、UTF-8 文本。播报完成回复为 0x41 0x4F，即 AO。vtx316_speak_wait 使用事件等待播报完成，超时 30000 ms。",
        "BLE Text 链路：_text_chr_access 先截获 CMD:，普通文本回调到 _ble_text_recv_handler，规范化 SAY:、全角冒号和空白后调用 vtx316_speak_wait。",
        "[此处插入 TTS 链路图]",
    ], [
        "#define VTX316_UART_NAME           \"uart1\"",
        "#define VTX316_UART_BAUDRATE       115200",
        "#define VTX316_FRAME_HEADER        0xFD",
        "#define VTX316_CMD_SPEAK           0x01",
        "#define VTX316_ENCODING_UTF8       0x05",
        "#define VTX316_REPLY_DONE_0        0x41",
        "#define VTX316_REPLY_DONE_1        0x4F",
        "vtx316_send_byte(VTX316_FRAME_HEADER);",
        "vtx316_send_byte((rt_uint8_t)(data_len >> 8));",
        "vtx316_send_byte((rt_uint8_t)(data_len & 0xFF));",
        "vtx316_send_byte(VTX316_CMD_SPEAK);",
        "vtx316_send_byte(VTX316_ENCODING_UTF8);",
    ])

    module(doc, "模块 7：右手音频采集与 AI 云识别模块（STT）", [
        "对应源文件：drv_sai_inmp441.c/.h、audio_capture_inmp441.c、audio_process.c/.h、voice_assistant.c、ai_cloud_service.c、web_client.c。",
        "INMP441 输出 24-bit I2S/PCM，工程目标为 16 kHz、16-bit、mono、little-endian PCM。SAI 注释记录 SCK=PA2、WS=PC0，SD 接线需以实际硬件为准。",
        "audio_process 使用能量阈值、VAD hangover 和最短语音帧数做端点检测；录音完成后 speech_data_handler 转为 PCM16 并释放 speech_ready_sem。",
        "voice_assistant 线程启动采集和 VAD，等待语音段，过滤过短音频，调用 ai_cloud_service_speech_to_text，再将识别文本显示到 OLED。讯飞模式通过 HTTP 代理发送 application/octet-stream 原始 PCM。",
        "[此处插入 STT 数据流图]",
    ], [
        "#define INMP441_SAMPLE_RATE         16000",
        "#define INMP441_BIT_WIDTH           24",
        "#define INMP441_CHANNEL_NUM         1",
        "#define SAI_DMA_BUFFER_SIZE         1024",
        "#define AUDIO_BUFFER_COUNT          4",
        "#define AUDIO_FRAME_SIZE            512",
        "#define VAD_THRESHOLD               20000",
        "#define VAD_HANGOVER_FRAMES         31",
        "#define VAD_MIN_SPEECH_FRAMES       10",
        "#define AUDIO_PROCESS_MAX_RECORD_SEC 2",
        "ret = ai_cloud_service_speech_to_text(audio_buffer, total_read, &ai_response);",
        "oled_show_stt_result(ai_response.text_result);",
    ])

    module(doc, "模块 8：OLED 状态显示模块", [
        "对应源文件：IIC/OLED/*.c、IIC/iic_thread.c、UART/uart_oled_thread.c。",
        "OLED 为 128×64 I2C 点阵屏，SCL/SDA 使用 PE1/PE2。IIC 线程优先级 20、栈 2048 字节，OLED 默认位于 TCA9548A 通道 3；若未探测到 TCA9548A，可直接挂接 I2C2。",
        "OLED 库提供初始化、清屏、刷新、字符串、数字、图形绘制等接口，字库包含 ASCII 与中文点阵。右手 STT 结果通过 oled_show_stt_result 显示。",
    ], [
        "#define OLED_TCA9548A_CHANNEL   3",
        "#define OLED_SCL_PIN    GET_PIN(E, 1)",
        "#define OLED_SDA_PIN    GET_PIN(E, 2)",
        "OLED_I2C_Init();",
        "tca9548a_init();",
        "if (tca9548a_is_present()) {",
        "    tca9548a_select_channel(OLED_TCA9548A_CHANNEL);",
        "}",
        "OLED_Init();",
        "OLED_Update();",
        "oled_show_system_status(\"OLED READY\");",
    ])

    module(doc, "模块 9：电池与电源监测模块", [
        "对应源文件：adc_battery.c/.h。",
        "ADC 使用 STM32 HAL 直接操作 ADC1 Channel 12，即 PC2 引脚。12 位 ADC、3.3 V 参考电压、分压比 2。battery 线程每 1 秒采样并进行 8 点滑动平均。",
        "电量按 3000-4200 mV 线性映射到 0-100%，对外提供 battery_get_voltage 与 battery_get_percentage，MSH 命令 bat 可查看当前电压和百分比。",
    ], [
        "#define BATTERY_THREAD_STACK_SIZE   1024",
        "#define BATTERY_THREAD_PRIORITY     22",
        "#define BATTERY_FULL_VOLTAGE        4200",
        "#define BATTERY_EMPTY_VOLTAGE       3000",
        "#define BATTERY_DIVIDER_RATIO       2",
        "#define ADC_REF_VOLTAGE             3300",
        "#define ADC_RESOLUTION              4096",
        "sConfig.Channel = ADC_CHANNEL_12;",
        "g_battery_voltage_mv = adc_to_battery_voltage(avg_adc);",
        "g_battery_percentage = calculate_percentage(g_battery_voltage_mv);",
    ])

    module(doc, "模块 10：辅助通信与调试模块", [
        "对应源文件：tcp_client.c/.h、nimble_hci_adapter.c、cyw43438_bt、rtdbg.h 使用点。",
        "TCP 默认服务器为 192.168.221.92:8266，线程名 tcp_cli，优先级 22、栈 4096、发送间隔 100 ms。MSH 导出 tcp_start/tcp_stop，左手还可暂存 translated_text 并随 TCP JSON 辅助上传。",
        "NimBLE HCI 适配器负责 HCI H4 TX/RX 和 mbuf 对接；CYW43438 UART7 与 WiFi SDIO 独立。日志通过 DBG_SECTION_NAME/DBG_TAG 与 DBG_LEVEL 控制。",
    ], [
        "#define TCP_SERVER_IP       \"192.168.221.92\"",
        "#define TCP_SERVER_PORT     8266",
        "#define TCP_THREAD_PRIORITY     22",
        "#define TCP_THREAD_STACK_SIZE   4096",
        "#define TCP_SEND_BUF_SIZE       4096",
        "#define TCP_RECV_BUF_SIZE       1024",
        "#define TCP_SEND_INTERVAL       100",
        "MSH_CMD_EXPORT_ALIAS(tcp_client_start, tcp_start, Start TCP client);",
        "MSH_CMD_EXPORT_ALIAS(tcp_client_stop, tcp_stop, Stop TCP client);",
        "struct os_mbuf *om = ble_hs_mbuf_from_flat(frag_buf, total_len);",
    ])

    module(doc, "模块 11：硬件接口与引脚分配表", [
        "I2C1：PB8-SCL、PB9-SDA，连接 TCA9548A #1 和 8 路 MPU6050。I2C2：PE1-SCL、PE2-SDA，连接 TCA9548A #2、2 路 MPU6050、ICM-20948 和 OLED。",
        "UART1：PF13-TX、PF12-RX，左手用于 VTX316，右手用于外部串口文本输入/OLED 显示。UART7：CYW43438 蓝牙 HCI H4。PC6：运行模式按键。PC2：ADC1 CH12 电池分压采样。PO5：LED 心跳。",
        "SAI/INMP441：SCK=PA2、WS=PC0，SD 引脚以最终硬件焊接为准。WiFi 使用 CYW43438 SDIO/WHD 资源。",
        "[此处插入 CYW43438 连接框图]",
    ], [
        "#define LED_PIN GET_PIN(O, 5)",
        "#define OP_MODE_BUTTON_PIN GET_PIN(C, 6)",
        "#define OLED_SCL_PIN GET_PIN(E, 1)",
        "#define OLED_SDA_PIN GET_PIN(E, 2)",
        "#define SAI2_SCK_PIN GPIO_PIN_2   /* PA2 */",
        "#define SAI2_FS_PIN  GPIO_PIN_0   /* PC0 */",
        "#define SAI2_SD_PIN  GPIO_PIN_3   /* PE3 */",
        "#define VTX316_UART_NAME \"uart1\"",
        "#define CYW43438_BT_UART_NAME \"uart7\"",
    ])

    module(doc, "模块 12：术语与参考文档", [
        "BLE/GATT/ATT/MTU：低功耗蓝牙属性通信相关术语；NimBLE：本工程使用的 BLE Host 协议栈；HCI：Host 与 Controller 之间的控制接口。",
        "IMU/MPU6050/ICM-20948/AK09916：惯性与磁力传感器术语；TTS/STT/VAD/IAT：语音合成、语音识别、端点检测与讯飞听写相关术语。",
        "RT-Thread/SCons/Keil MDK：运行时系统与构建工具。开源组件许可证包括 RT-Thread Apache 2.0、CMSIS Apache 2.0、STM32 HAL BSD 3-Clause、NimBLE Apache 2.0。",
    ])

    doc.add_heading("附录 A：测试与验证", 1)
    lines(doc, [
        "1. 上电后确认 PO5 心跳、串口日志、IIC/OLED、MPU、battery 线程启动。",
        "2. 使用 BLE 工具扫描 ART-Pi2-IMU-L/R，连接并发现 A74D0001 服务。",
        "3. 订阅 Notify 后发送 CMD:START，确认 [FRAG] 分片按 90 ms 周期输出；发送 CMD:STOP 后停止。",
        "4. 断开 BLE 后确认设备重新广播；重连订阅后恢复发送。",
        "5. 模拟 IMU 通道异常，确认 CSV 字段数不变且对应通道补 0。",
        "6. 左手写入 SAY:测试，确认 VTX316 播报；右手录音后确认 STT 文本显示到 OLED。",
        "7. 执行 bat、tcp_start、tcp_stop、vtx316_say 等 MSH 命令验证调试接口。",
    ])

    doc.add_heading("附录 B：配置与部署", 1)
    lines(doc, [
        "构建：进入 art_pi2_left 或 art_pi2_right，使用 RT-Thread ENV/SCons 生成或更新 Keil 工程；rtconfig.py 可通过 RTT_CC=keil 切换 armclang。",
        "烧录：使用 ST-Link/J-Link 下载固件到 ART-Pi2。串口：连接控制台查看 rt_kprintf/rtdbg 日志。",
        "运行配置：确认 WHD 固件/CLM/NVRAM 资源可访问，WiFi SSID/口令、TCP IP、讯飞代理 URL 已按现场网络调整。",
        "提交前处理：voice_assistant_config.h 中 APPID/APIKey/APISecret、WiFi 密码等敏感字段应替换为占位符或脱敏版本。",
    ])

    doc.add_heading("待确认事项", 1)
    lines(doc, [
        "[待确认：右手 INMP441 实际 SD 引脚以硬件焊接为准；源码注释中存在 PE3 与 PE7 两种说明。]",
        "[待确认：讯飞 WebSocket/HMAC/IAT 细节当前主要位于 PC 代理 tools/xfyun_proxy.py；若申报固件直接实现，应补充 MCU 端源码。]",
        "[待确认：BLE MAC 地址随开发板芯片固件变化，批量更换硬件时应更新实测 MAC 或删除固定 MAC 描述。]",
        "[待确认：ROCK 5B 上位机识别程序不在本次申报范围内，提交材料时避免混入其源码页。]",
    ])

    doc.save(OUT)


if __name__ == "__main__":
    main()
